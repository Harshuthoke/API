from flask import Flask, render_template, request, redirect, send_file
from flask_pymongo import PyMongo
from bson.objectid import ObjectId
from docx import Document
import io

app = Flask(__name__)

# MongoDB Atlas connection string
app.config["MONGO_URI"] = "mongodb+srv://harshal:Harshal2022@cluster0.u5i2m.mongodb.net/form?retryWrites=true&w=majority&appName=Cluster0"
mongo = PyMongo(app)

# Route for rendering the form
@app.route('/')
def index():
    return render_template('form.html')

# Route for handling form submissions and generating the report document
@app.route('/submit', methods=['POST'])
def submit():
    if request.method == 'POST':
        # Extract form data
        form_data = {
            "case_number": request.form.get('case_number'),
            "date_of_report": request.form.get('date_of_report'),
            "report_prepared_by": request.form.get('report_prepared_by'),
            "model": request.form.get('model'),
            "color": request.form.get('color'),
            "safety_glass": request.form.get('safety_glass'),
            "back_cover": request.form.get('back_cover'),
            "ram": request.form.get('ram'),
            "internal_memory": request.form.get('internal_memory'),
            "camera_specs": request.form.get('camera_specs'),
            "cameras_check": request.form.get('cameras_check'),
            "battery_percentage": request.form.get('battery_percentage'),
            "sim_slots": request.form.get('sim_slots'),
            "sim_provider": request.form.get('sim_provider'),
            "wifi_connected": request.form.get('wifi_connected'),
            "bluetooth_status": request.form.get('bluetooth_status'),
            "sd_card": request.form.get('sd_card'),
            "sd_capacity": request.form.get('sd_capacity'),
            "mobile_uptime": request.form.get('mobile_uptime'),
            "time_zone": request.form.get('time_zone'),
            "language": request.form.get('language'),
            "installed_apps": request.form.get('installed_apps'),
            "geo_location": request.form.get('geo_location'),
            "build_no": request.form.get('build_no'),
            "kernel_version": request.form.get('kernel_version'),
            "iccid": request.form.get('iccid'),
            "imsi": request.form.get('imsi'),
            "meid": request.form.get('meid'),
            "airplane_mode": request.form.get('airplane_mode')
        }
        
        # Insert data into MongoDB
        inserted_id = mongo.db.device_info.insert_one(form_data).inserted_id

        # Load the report template document
        document = Document('./Report Format.docx')

        # Replace placeholders in the document with actual user data
        for para in document.paragraphs:
            if '[Insert Case Number]' in para.text:
                para.text = para.text.replace('[Insert Case Number]', form_data['case_number'])
            if '[Insert Date]' in para.text:
                para.text = para.text.replace('[Insert Date]', form_data['date_of_report'])
            if '[Insert Name and Title]' in para.text:
                para.text = para.text.replace('[Insert Name and Title]', form_data['report_prepared_by'])
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if '[Insert Device Model]' in cell.text:
                        cell.text = cell.text.replace('[Insert Device Model]', form_data['model'])
                    if '[Insert Color]' in cell.text:
                        cell.text = cell.text.replace('[Insert Color]', form_data['color'])
                    if '[Safety Glass Yes or No]' in cell.text:
                        cell.text = cell.text.replace('[Safety Glass Yes or No]', form_data['safety_glass'])
                    if '[Back Cover Yes or No]' in cell.text:
                        cell.text = cell.text.replace('[Back Cover Yes or No]', form_data['back_cover'])
                    if '[Insert RAM]' in cell.text:
                        cell.text = cell.text.replace('[Insert RAM]', form_data['ram'])
                    if '[Insert Internal Memory]' in cell.text:
                        cell.text = cell.text.replace('[Insert Internal Memory]', form_data['internal_memory'])
                    if '[Insert Camera Details]' in cell.text:
                        cell.text = cell.text.replace('[Insert Camera Details]', form_data['camera_specs'])
                    if '[Insert Cameras Check]' in cell.text:
                        cell.text = cell.text.replace('[Insert Cameras Check]', form_data['cameras_check'])
                    if '[Insert Battery Percentage]' in cell.text:
                        cell.text = cell.text.replace('[Insert Battery Percentage]', form_data['battery_percentage'])
                    if '[Insert SIM Slots]' in cell.text:
                        cell.text = cell.text.replace('[Insert SIM Slots]', form_data['sim_slots'])
                    if '[Insert SIM Provider]' in cell.text:
                        cell.text = cell.text.replace('[Insert SIM Provider]', form_data['sim_provider'])
                    if '[Insert Wi-Fi Status]' in cell.text:
                        cell.text = cell.text.replace('[Insert Wi-Fi Status]', form_data['wifi_connected'])
                    if '[Insert Bluetooth Status]' in cell.text:
                        cell.text = cell.text.replace('[Insert Bluetooth Status]', form_data['bluetooth_status'])
                    if '[Insert SD Card Present]' in cell.text:
                        cell.text = cell.text.replace('[Insert SD Card Present]', form_data['sd_card'])
                    if '[Insert SD Capacity]' in cell.text:
                        cell.text = cell.text.replace('[Insert SD Capacity]', form_data['sd_capacity'])
                    if '[Insert Mobile Uptime]' in cell.text:
                        cell.text = cell.text.replace('[Insert Mobile Uptime]', form_data['mobile_uptime'])
                    if '[Insert Time Zone]' in cell.text:
                        cell.text = cell.text.replace('[Insert Time Zone]', form_data['time_zone'])
                    if '[Insert Language]' in cell.text:
                        cell.text = cell.text.replace('[Insert Language]', form_data['language'])
                    if '[Insert Installed Apps]' in cell.text:
                        cell.text = cell.text.replace('[Insert Installed Apps]', form_data['installed_apps'])
                    if '[Insert Geo Location]' in cell.text:
                        cell.text = cell.text.replace('[Insert Geo Location]', form_data['geo_location'])
                    if '[Insert Build Number]' in cell.text:
                        cell.text = cell.text.replace('[Insert Build Number]', form_data['build_no'])
                    if '[Insert Kernel Version]' in cell.text:
                        cell.text = cell.text.replace('[Insert Kernel Version]', form_data['kernel_version'])
                    if '[Insert ICCID]' in cell.text:
                        cell.text = cell.text.replace('[Insert ICCID]', form_data['iccid'])
                    if '[Insert IMSI]' in cell.text:
                        cell.text = cell.text.replace('[Insert IMSI]', form_data['imsi'])
                    if '[Insert MEID]' in cell.text:
                        cell.text = cell.text.replace('[Insert MEID]', form_data['meid'])
                    if '[Insert Airplane Mode]' in cell.text:
                        cell.text = cell.text.replace('[Insert Airplane Mode]', form_data['airplane_mode'])


        # Save the modified document to an in-memory stream
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)

        # Serve the generated document as a downloadable file
        return send_file(buffer, as_attachment=True, download_name=f'report_{form_data["case_number"]}.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

if __name__ == '__main__':
    app.run(debug=True)
